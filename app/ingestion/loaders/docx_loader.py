from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.document import Document
from app.ingestion.loaders.base_loader import BaseLoader


class DOCXLoader(BaseLoader):

    def load(self, file_path: str) -> Document:

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        metadata = {
            "filename": Path(file_path).name,
            "filetype": "docx",
        }

        return Document(
            content=text,
            metadata=metadata,
        )